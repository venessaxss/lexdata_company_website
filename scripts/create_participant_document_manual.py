from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_FILE = OUT_DIR / "Participant-Certificate-and-Receipt-Manual.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_BLUE = "F4F7FB"
GOLD = "B08D39"
PALE_GOLD = "FFF8E8"
GREEN = "177245"
PALE_GREEN = "EAF7F0"
RED = "9B1C1C"
PALE_RED = "FDEEEE"
INK = "1F2937"
MUTED = "5F6B7A"
WHITE = "FFFFFF"
GRID = "CFD8E3"


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_border(paragraph, color=GRID, size=8, space=3, side="bottom"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_numbering_definitions(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def make_num(abstract_id, num_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    make_num(next_abs, next_num, "decimal", "%1.")
    make_num(next_abs + 1, next_num + 1, "bullet", "\u2022", "Symbol")
    return next_num, next_num + 1


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def add_list_item(doc, text, num_id, bold_lead=None):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label}: ")
    set_font(r1, bold=True, color=accent)
    r2 = p.add_run(text)
    set_font(r2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_step(doc, number, title, body):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [900, 8460])
    left, right = table.rows[0].cells
    set_cell_shading(left, NAVY)
    set_cell_shading(right, LIGHT_BLUE)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(number))
    set_font(r, size=18, bold=True, color=WHITE)
    p2 = right.paragraphs[0]
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(title)
    set_font(r2, size=11.5, bold=True, color=NAVY)
    p3 = right.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.line_spacing = 1.15
    r3 = p3.add_run(body)
    set_font(r3, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_route(doc, label, route):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(f"{label}: ")
    set_font(r1, bold=True, color=NAVY)
    r2 = p.add_run(route)
    set_font(r2, name="Consolas", size=10, color=BLUE, bold=True)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run("LEXDATA")
    set_font(r1, size=9, bold=True, color=NAVY)
    r2 = p.add_run("   |   Participant Document Guide")
    set_font(r2, size=9, color=MUTED)
    set_paragraph_border(p, color=GRID, size=6, space=3)

    footer = section.footer
    p2 = footer.paragraphs[0]
    set_paragraph_border(p2, color=GRID, size=4, space=3, side="top")
    add_page_number(p2)


def build_manual():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        configure_header_footer(section)
    decimal_num, bullet_num = add_numbering_definitions(doc)

    # Editorial cover (named override for title size and cover whitespace).
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PARTICIPANT GUIDE")
    set_font(r, size=11, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Certificates &\nPayment Receipts")
    set_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("How to prepare your preferred name, qualify, receive, download, and verify your LexData documents")
    set_font(r, size=14, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run("For course and workshop participants")
    set_font(r, size=11, italic=True, color=MUTED)

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    for cell, label, value in zip(
        table.rows[0].cells,
        ["RECEIPTS", "CERTIFICATES", "VERIFICATION"],
        ["After confirmed payment", "After completion + approval", "Public status check"],
    ):
        set_cell_shading(cell, LIGHT_BLUE)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        a = p1.add_run(label)
        set_font(a, size=9, bold=True, color=BLUE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        b = p2.add_run(value)
        set_font(b, size=10.5, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("Version 1.0 | August 2026")
    set_font(r, size=9.5, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "1. Start here", 1)
    add_callout(
        doc,
        "Important",
        "Workshop certificates require an application. After the workshop is marked completed, open Certificates & Receipts, apply for that specific workshop, and wait for admin approval.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    add_heading(doc, "The complete participant journey", 2)
    add_step(doc, 1, "Set your preferred printed name", "Open your profile and enter your name exactly as it should appear. You may use English, Chinese, Arabic, Urdu, or another supported script.")
    add_step(doc, 2, "Register and complete payment", "Enroll in the course or workshop. Follow the payment instructions and upload proof if a manual payment method is used.")
    add_step(doc, 3, "Complete the learning requirement", "Complete every published course lesson, or attend the workshop and wait for the admin to mark your participation completed.")
    add_step(doc, 4, "Apply for the workshop certificate", "Open Certificates & Receipts, select the completed workshop, confirm the exact printed name, and submit the application.")
    add_step(doc, 5, "Open your released documents", "After admin approval, open the issued certificate, then print it or save it as PDF.")

    add_heading(doc, "Before you begin", 2)
    for item in [
        "Use your own participant account; documents are linked to that account.",
        "Make sure your email address is correct and accessible.",
        "Enter your preferred name before completion or payment approval.",
        "Keep your payment reference or transfer evidence until the receipt is released.",
    ]:
        add_list_item(doc, item, bullet_num)

    add_route(doc, "Profile", "Dashboard > My Profile")
    add_route(doc, "Documents", "Dashboard > Certificates & Receipts")
    add_route(doc, "Support", "Dashboard > Messages")
    doc.add_page_break()

    add_heading(doc, "2. Set the name printed on your documents", 1)
    add_body(doc, "Your preferred printed name is used for both certificates and payment receipts. The system stores a snapshot of the name when a document is prepared or issued, so later profile edits do not silently change an already released document.")
    add_heading(doc, "How to set your name", 2)
    steps = [
        "Sign in to your LexData account.",
        "Open Dashboard, then select My Profile.",
        "Find Preferred name on certificates and receipts.",
        "Enter the exact spelling, capitalization, spacing, and script you want printed.",
        "Select Save Profile.",
    ]
    for item in steps:
        add_list_item(doc, item, decimal_num)

    add_heading(doc, "Name examples", 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    rows = [
        ("English", "Dr. Ayesha Khan"),
        ("Chinese", "王晓明"),
        ("Arabic", "د. أحمد العتيبي"),
        ("Urdu", "عائشہ خان"),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    table.rows[0].cells[0].text = "Script"
    table.rows[0].cells[1].text = "Example preferred name"
    set_table_geometry(table, [2700, 6660])
    mark_header_row(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, color=NAVY)
    for row in table.rows[1:]:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                set_font(run)

    add_callout(doc, "Check carefully", "If a certificate is already issued with the wrong name, changing your profile will not change that certificate. Contact an admin through Messages and explain the correction needed.", fill=PALE_RED, accent=RED)
    doc.add_page_break()

    add_heading(doc, "3. How to receive a payment receipt", 1)
    add_body(doc, "A payment upload is evidence for review; it is not the official payment receipt. The platform releases the official organization receipt only after the payment amount and currency are confirmed.")
    add_heading(doc, "Online card or Stripe payment", 2)
    for item in [
        "Select the course or workshop and continue to checkout.",
        "Complete payment on the secure checkout page.",
        "Wait for Stripe to confirm the exact amount and currency.",
        "When confirmation succeeds, the receipt is generated automatically and appears in Certificates & Receipts.",
    ]:
        add_list_item(doc, item, decimal_num)

    add_heading(doc, "Manual bank transfer or local payment", 2)
    for item in [
        "Register for the workshop and read the payment instructions sent by the LexData team.",
        "Complete the transfer using the instructed account and amount.",
        "Open the workshop page and upload your payment evidence when the upload option is available.",
        "An admin reviews the evidence, records the received amount and currency, and confirms payment.",
        "After confirmation, access is unlocked and the official payment receipt is released automatically.",
    ]:
        add_list_item(doc, item, decimal_num)

    add_callout(doc, "Receipt gate", "A receipt cannot be generated for a pending, unverified, zero, failed, cancelled, or refunded payment. Free or waived participation does not create a paid receipt.", fill=PALE_GOLD, accent=GOLD)

    add_heading(doc, "Receipt jurisdiction", 2)
    add_body(doc, "The receipt jurisdiction is selected by the issuing LexData entity, not by your nationality or residence. Your receipt may therefore show Pakistan, Saudi Arabia, or China according to the organization that received the payment.")
    add_body(doc, "The platform document is proof of payment by default. It is not automatically an FBR fiscal invoice, a ZATCA FATOORAH tax invoice, or an official Chinese fapiao. If you need a tax document, contact the admin and provide the buyer or organization details required for external tax-authority issuance.")
    doc.add_page_break()

    add_heading(doc, "4. How to qualify for a certificate", 1)
    add_heading(doc, "Course certificate", 2)
    add_body(doc, "For a course certificate, the platform prepares your certificate draft when all of the following conditions are satisfied:")
    for item in [
        "Your course enrollment is approved or confirmed.",
        "Your payment status is paid or waived.",
        "Every published lesson in the course is marked complete.",
        "Your account has a preferred certificate name.",
    ]:
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "Workshop certificate", 2)
    add_body(doc, "Attend the workshop and meet the organizer's participation or completion requirements. After an admin marks your registration completed, an application form appears under Certificates & Receipts. Select that workshop, confirm your preferred printed name, add an optional note, and submit the application.")

    add_heading(doc, "Admin review and release", 2)
    add_step(doc, 1, "Application submitted", "The application status becomes Pending. It is tied to your completed registration for that specific workshop.")
    add_step(doc, 2, "Admin checks the record", "The admin confirms completion, participant identity, preferred name, workshop, issuing jurisdiction, and active certificate template.")
    add_step(doc, 3, "Certificate generated from the template", "Approval places your name and workshop details onto the admin-uploaded template and releases the certificate as Issued.")

    add_callout(doc, "Application required", "Workshop completion makes you eligible, but it does not issue the certificate by itself. Submit the application for that workshop. If it remains pending after the normal review period, contact support through Dashboard > Messages.", fill=PALE_GREEN, accent=GREEN)
    doc.add_page_break()

    add_heading(doc, "5. Open, download, print, and verify", 1)
    add_heading(doc, "Open your document center", 2)
    for item in [
        "Sign in and open Dashboard.",
        "Select Certificates & Receipts.",
        "Find the document card and check its status.",
        "Select Open document when the status is Issued.",
        "Select Print / save as PDF in the document view.",
    ]:
        add_list_item(doc, item, decimal_num)

    add_heading(doc, "Understand document statuses", 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    status_rows = [
        ("Pending review", "Certificate prepared but awaiting admin approval; printing is locked."),
        ("Issued", "Valid and available to open, print, save, and verify."),
        ("Revoked", "Certificate is no longer valid; the verification page displays the current status."),
        ("Void", "Receipt is no longer valid, commonly because the underlying payment was refunded or cancelled."),
    ]
    for status, meaning in status_rows:
        cells = table.add_row().cells
        cells[0].text = status
        cells[1].text = meaning
    table.rows[0].cells[0].text = "Status"
    table.rows[0].cells[1].text = "What it means"
    set_table_geometry(table, [2700, 6660])
    mark_header_row(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, color=NAVY)
    for row in table.rows[1:]:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10.5)

    add_heading(doc, "Verify authenticity", 2)
    add_body(doc, "Every issued certificate and receipt includes a document number, a verification code, and a public verification address. A university, employer, sponsor, or payer can open that address to check the participant name, program, jurisdiction, issue date, amount for receipts, and current validity status.")
    add_callout(doc, "Share safely", "Share the public verification link when a third party needs to confirm authenticity. Do not share your password, payment card details, private transfer records, or dashboard access.", fill=LIGHT_BLUE, accent=BLUE)
    doc.add_page_break()

    add_heading(doc, "6. Troubleshooting and support", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    issues = [
        ("Receipt not visible", "Check that payment is confirmed, not merely uploaded or under review. Then refresh Certificates & Receipts."),
        ("Confirmed amount is wrong", "Do not use the receipt. Contact the admin through Messages and include the payment reference and expected amount."),
        ("Certificate still pending", "Confirm every lesson is complete, or wait for workshop attendance to be marked completed. Then contact support if needed."),
        ("Name is wrong before issue", "Update Preferred name on certificates and receipts in My Profile before admin approval."),
        ("Name is wrong after issue", "Contact the admin. Issued documents keep a permanent snapshot and require a controlled correction."),
        ("Document shows revoked or void", "Read the verification status and contact support for the recorded reason or next steps."),
        ("Need a tax invoice", "Request it from the admin. The platform payment receipt is not automatically an FBR invoice, ZATCA invoice, or fapiao."),
        ("Cannot print", "Open the issued document in a modern browser and use Print / save as PDF. Allow the print dialog if the browser blocks it."),
    ]
    for issue, action in issues:
        cells = table.add_row().cells
        cells[0].text = issue
        cells[1].text = action
    table.rows[0].cells[0].text = "Issue"
    table.rows[0].cells[1].text = "What to do"
    set_table_geometry(table, [3000, 6360])
    mark_header_row(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, color=NAVY)
    for row in table.rows[1:]:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10.25)

    add_heading(doc, "What to include in a support message", 2)
    for item in [
        "Your registered email address.",
        "Course or workshop title.",
        "Document number, if one is displayed.",
        "Payment reference and amount for receipt questions.",
        "A short description of the problem; do not send passwords or complete card details.",
    ]:
        add_list_item(doc, item, bullet_num)
    add_route(doc, "Contact route", "Dashboard > Messages")

    add_callout(doc, "Final reminder", "Set your preferred name early, keep payment evidence until confirmation, complete every learning requirement, and use the public verification link whenever a third party needs to validate your document.", fill=PALE_GREEN, accent=GREEN)

    # Document core properties.
    props = doc.core_properties
    props.title = "LexData Participant Certificate and Receipt Manual"
    props.subject = "Participant instructions for certificate and payment receipt issuance"
    props.author = "LexData"
    props.keywords = "LexData, certificate, receipt, participant manual, verification"

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_manual())
